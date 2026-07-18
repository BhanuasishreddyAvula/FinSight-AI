"""
docx_parser.py — Structured DOCX extraction using python-docx.

Extracts heading hierarchy, paragraphs, lists, and tables, converting
each into a DocBlock for the unified downstream pipeline.

Heading levels are detected from the Word paragraph style name
(e.g. "Heading 1", "Heading 2", "Title").
"""
from __future__ import annotations
from docx import Document
from docx.oxml.ns import qn
from ingestion.models import DocBlock


# Map Word style names to heading levels
_HEADING_STYLES: dict[str, int] = {
    "title": 1,
    "heading 1": 1,
    "heading 2": 2,
    "heading 3": 3,
    "heading 4": 4,
    "heading 5": 5,
    "heading 6": 6,
}


def parse_docx(file_path: str) -> list[DocBlock]:
    """
    Parse a DOCX file and return a list of structured DocBlocks.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        List of DocBlock objects with heading hierarchy preserved.
    """
    doc = Document(file_path)
    blocks: list[DocBlock] = []
    block_index = 0
    section_stack: list[tuple[int, str]] = []  # (level, text)

    # Iterate document body elements in reading order
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""
        heading_level = _HEADING_STYLES.get(style_name, 0)

        if heading_level > 0:
            # Update section breadcrumb
            while section_stack and section_stack[-1][0] >= heading_level:
                section_stack.pop()

            current_path = [s[1] for s in section_stack]
            section_stack.append((heading_level, text))

            blocks.append(DocBlock(
                block_type="heading",
                text=text,
                level=heading_level,
                page_number=None,
                section_path=current_path,
                block_index=block_index,
            ))
        else:
            # Detect list items by paragraph numbering / bullet format
            is_list = _is_list_item(para)
            block_type = "list" if is_list else "paragraph"
            current_path = [s[1] for s in section_stack]

            blocks.append(DocBlock(
                block_type=block_type,
                text=text,
                level=0,
                page_number=None,
                section_path=current_path,
                block_index=block_index,
            ))

        block_index += 1

    # Extract tables — each row becomes part of the table block
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))

        if rows:
            # Use the section path at this point (approximate — DOCX has no table position)
            current_path = [s[1] for s in section_stack]
            blocks.append(DocBlock(
                block_type="table",
                text="\n".join(rows),
                level=0,
                page_number=None,
                section_path=current_path,
                block_index=block_index,
            ))
            block_index += 1

    return blocks


def _is_list_item(para) -> bool:
    """
    Detect if a paragraph is a bullet or numbered list item.
    Checks for numPr XML element (Word's numbering property).
    """
    try:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                return True
    except Exception:
        pass
    # Fallback: style name contains "list"
    style_name = para.style.name.lower() if para.style else ""
    return "list" in style_name
